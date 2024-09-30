from flask import Flask, request, jsonify, render_template, send_from_directory
from werkzeug.utils import secure_filename
import os
import PyPDF2
import requests
import json

from google.cloud import aiplatform
from google.oauth2 import service_account


from langchain_google_vertexai import VertexAI
from langchain_core.runnables import RunnableLambda
from langchain_core.messages import AIMessage
from langchain.prompts import PromptTemplate
from langchain import LLMChain
from langchain.memory import ConversationBufferMemory



import re
import subprocess
import docx
from PyPDF2 import PdfReader
from transformers import pipeline
from bs4 import BeautifulSoup




app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


MODEL_ID = "gemini-pro"  # Replace with your actual endpoint ID
#PROJECT_ID = 'flexitext-435112'
PROJECT_ID = "grantgpt-433615"
REGION = "us-central1"


# Ensure the uploads directory exists
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Initialize memory to remember past conversations (or use ConversationSummaryMemory)
memory = ConversationBufferMemory()


#Code to authenticate with the Google Cloud Platform
def authicate_google():
    
    key_path = "C:/Users/a884470/backend/grantgp.json"
    
    credentials = service_account.Credentials.from_service_account_file(key_path)


    aiplatform.init(
        project=PROJECT_ID,
        credentials=credentials,
        location=REGION
    )
    print("Authenticated with Google Cloud Platform")
    

#TO DO: Improve the function to handle more markdown formatting
#Make sure to handle undelining and headings
#There must be no color change

def add_markdown_to_word(doc, markdown_text):
    # Split the markdown text into lines
    lines = markdown_text.splitlines()

    for line in lines:
        # Handle headers (Headings level 1-6)
        header_match = re.match(r'^(#{1,6})\s*(.*)', line)
        if header_match:
            header_level = len(header_match.group(1))
            header_text = header_match.group(2)
            # Add header text to the Word document with the appropriate header level
            doc.add_heading(header_text, level=header_level)
            continue

        # Create a new paragraph
        para = doc.add_paragraph()

        # Handle bold (**bold text**)
        line = re.sub(r'\*\*(.*?)\*\*', r'{BOLD:\1}', line)

        # Handle italic (*italic text* or _italic text_)
        line = re.sub(r'(\*|_)(.*?)\1', r'{ITALIC:\2}', line)

        # Handle underline (__underline__)
        line = re.sub(r'__(.*?)__', r'{UNDERLINE:\1}', line)

        # Split the line based on our temporary placeholders (BOLD, ITALIC, UNDERLINE)
        parts = re.split(r'(\{BOLD:.*?\}|\{ITALIC:.*?\}|\{UNDERLINE:.*?\})', line)

        for part in parts:
            if part.startswith("{BOLD:") and part.endswith("}"):
                bold_text = part[6:-1]
                run = para.add_run(bold_text)
                run.bold = True
            elif part.startswith("{ITALIC:") and part.endswith("}"):
                italic_text = part[8:-1]
                run = para.add_run(italic_text)
                run.italic = True
            elif part.startswith("{UNDERLINE:") and part.endswith("}"):
                underline_text = part[11:-1]
                run = para.add_run(underline_text)
                run.underline = True
            else:
                para.add_run(part)

def extract_text(filepath, filename):
    """
    Extracts text from a .docx or .pdf file.
    
    Args:
        filepath (str): Full path to the file.
        filename (str): Name of the file (to detect file type).
    
    Returns:
        str: Extracted text.
    """
    
    # Extract text from a .docx file
    if filename.endswith('.docx'):
        doc = docx.Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs])
    
    # Extract text from a .pdf file
    elif filename.endswith('.pdf'):
        reader = PdfReader(open(filepath, 'rb'))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    else:
        raise ValueError("Unsupported file format. Please provide a .docx or .pdf file.")

def chunk_text(text, max_length=4000, output_dir="chunks"):
    """
    Split the text into chunks of specified maximum length without cutting in the middle of sentences,
    and save each chunk as a separate .txt file in the specified directory.

    Args:
        text (str): The input text to be chunked.
        max_length (int): The maximum length of each chunk.
        output_dir (str): Directory where the chunked .txt files will be saved.
    
    Returns:
        list: A list of chunked text parts.
    """
    
    # Ensure the output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Split the text into paragraphs
    paragraphs = text.split("\n")
    
    chunks = []
    current_chunk = ""
    chunk_index = 1
    
    # Build chunks without exceeding max_length
    for para in paragraphs:
        # If adding the current paragraph exceeds max_length, save the current chunk
        if len(current_chunk) + len(para) + 1 > max_length:  # +1 for newline
            chunks.append(current_chunk.strip())
            
            # Save the current chunk to a file
            chunk_filename = os.path.join(output_dir, f"chunk_{chunk_index}.txt")
            with open(chunk_filename, 'w', encoding='utf-8') as f:
                f.write(current_chunk.strip())
            print(f"Chunk {chunk_index} saved to {chunk_filename}")
            
            # Prepare for the next chunk
            current_chunk = para + "\n"
            chunk_index += 1
        else:
            current_chunk += para + "\n"
    
    # Add and save the last chunk if it exists
    if current_chunk:
        chunks.append(current_chunk.strip())
        chunk_filename = os.path.join(output_dir, f"chunk_{chunk_index}.txt")
        with open(chunk_filename, 'w', encoding='utf-8') as f:
            f.write(current_chunk.strip())
        print(f"Chunk {chunk_index} saved to {chunk_filename}")
    
    print(f"Text chunked into {len(chunks)} chunks")
    return chunks

def process_long_text(user_prompt, text):
    authicate_google()
    # Function to chunk the text
    chunks = chunk_text(text)
    processed_chunks = []
    
    
    for chunk_index, chunk in enumerate(chunks):
        # Prepare the prompt for the Vertex AI call
        prompt = f"""
        RULES:
        1) Follow the user prompt to modify the text as needed. 
        2) Do not change the user input.
        3) Do NOT remove any information from the text.
        4) Make sure the text is clear and easy to read.
        5) Make sure the text is grammatically correct.
        6) Do NOT add any additional information and/or text unless requested by the user.
        7) Do NOT add any labels, titles, or headings to the text.
        8) Do NOT add any additional formatting to the text.
        9) Do NOT forget that the text could be a part of a larger document, so make sure it fits in with the rest of the text.
        User input: {user_prompt}
        Text: {chunk}
        """
        print(f"Prompt: Created")

        # Fallback response for error handling
        empty_response = RunnableLambda(
            lambda x: AIMessage(content="Error processing document")
        )

            # Initialize the Vertex AI model client
        vertex_ai_client = VertexAI(
        temperature=0, model_name="gemini-1.0-pro", max_tokens=1024
        ).with_fallbacks([empty_response])

        # Define the input for the LLMChain
        var = {
            "text": chunk
        }

        # Create and run the LLMChain
        chain = LLMChain(
            llm=vertex_ai_client,
            prompt=PromptTemplate(
                input_variables=["text"],
                template=prompt
            )
        )
        response = chain.run(var)
        # Save each response as a separate file in the chunks directory
        response_filename = os.path.join("chunks", f"response_{chunk_index}.txt")
        with open(response_filename, 'w', encoding='utf-8') as f:
            f.write(response)
        print(f"Response {chunk_index} saved to {response_filename}")
        processed_chunks.append(response)

    # Combine all processed chunks back into a single string
    processed_chunks = "\n".join(processed_chunks)

    return processed_chunks

    



@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    user_prompt = request.form.get('prompt')

    
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    # Extract text and call Vertex AI
    content = extract_text(filepath, filename)
    print(f"Extracted text: {content}...")


    # Process the text using Vertex AI
    processed_text = process_long_text(user_prompt,content)
    


    #Create and save the Word document
    word_filename = f"processed_{filename.replace('.pdf', '.docx')}"
    word_filepath = os.path.join(app.config['UPLOAD_FOLDER'], word_filename)
    
    new_doc = docx.Document()
    add_markdown_to_word(new_doc, processed_text)  # Format the processed text
    new_doc.save(word_filepath)
    
    # Convert the Word document to PDF using LibreOffice in headless mode
    pdf_filename = f"processed_{filename.replace('.pdf', '.pdf')}"
    pdf_filepath = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
    libreoffice_path = r'C:\Program Files\LibreOffice\program\soffice.exe'

    try:
        result = subprocess.run([
            libreoffice_path, 
            '--headless', 
            '--convert-to', 
            'pdf', 
            word_filepath, 
            '--outdir', 
            UPLOAD_FOLDER
        ], check=True, capture_output=True, text=True)
        print("LibreOffice conversion output:", result.stdout)
        print("LibreOffice conversion error output:", result.stderr)
    except subprocess.CalledProcessError as e:
        print("An error occurred while converting the document to PDF.")
        print(f"Error message: {e.stderr}")
        
    
    # Determine the correct file URL to return
    if filename.endswith('.pdf'):
        file_url = f'/download/{pdf_filename}'
    else:
        file_url = f'/download/{word_filename}'

    return jsonify({'file_url': file_url}), 200


@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
